#!/usr/bin/env python3
"""
File Parser Module
Handles parsing of various document types for the NZ Legal RAG system
"""

import io
import os
import re
from pathlib import Path
from typing import Dict, List, Tuple, Optional, BinaryIO
from dataclasses import dataclass
import tempfile
import zipfile

@dataclass
class ParsedDocument:
    """Represents a parsed document"""
    filename: str
    content: str
    metadata: Dict
    file_type: str
    pages: Optional[int] = None
    sheets: Optional[int] = None


class FileParser:
    """Parser for various document types"""
    
    SUPPORTED_TYPES = {
        '.txt': 'text/plain',
        '.pdf': 'application/pdf',
        '.doc': 'application/msword',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        '.xls': 'application/vnd.ms-excel',
        '.html': 'text/html',
        '.htm': 'text/html',
        '.md': 'text/markdown',
        '.csv': 'text/csv',
        '.json': 'application/json',
    }
    
    def __init__(self):
        self.parsers = {
            '.txt': self._parse_text,
            '.md': self._parse_text,
            '.csv': self._parse_text,
            '.json': self._parse_text,
            '.html': self._parse_html,
            '.htm': self._parse_html,
            '.pdf': self._parse_pdf,
            '.doc': self._parse_doc,
            '.docx': self._parse_docx,
            '.xlsx': self._parse_xlsx,
            '.xls': self._parse_xlsx,
        }
    
    def parse_file(self, file_content: bytes, filename: str) -> ParsedDocument:
        """
        Parse a file and return structured content
        
        Args:
            file_content: Raw bytes of the file
            filename: Original filename
            
        Returns:
            ParsedDocument with extracted content and metadata
        """
        ext = Path(filename).suffix.lower()
        
        if ext not in self.SUPPORTED_TYPES:
            raise ValueError(f"Unsupported file type: {ext}. Supported: {list(self.SUPPORTED_TYPES.keys())}")
        
        parser = self.parsers.get(ext, self._parse_text)
        content, metadata = parser(file_content, filename)
        
        return ParsedDocument(
            filename=filename,
            content=content,
            metadata=metadata,
            file_type=ext
        )
    
    def parse_multiple(self, files: List[Tuple[bytes, str]]) -> List[ParsedDocument]:
        """Parse multiple files"""
        results = []
        errors = []
        
        for file_content, filename in files:
            try:
                doc = self.parse_file(file_content, filename)
                results.append(doc)
            except Exception as e:
                errors.append({"filename": filename, "error": str(e)})
        
        return results, errors
    
    def parse_zip(self, zip_content: bytes) -> Tuple[List[ParsedDocument], List[Dict]]:
        """Parse a zip file containing multiple documents"""
        results = []
        errors = []
        
        try:
            with zipfile.ZipFile(io.BytesIO(zip_content)) as zf:
                for info in zf.infolist():
                    # Skip directories and hidden files
                    if info.filename.endswith('/') or info.filename.startswith('__'):
                        continue
                    
                    # Skip macOS metadata files
                    if info.filename.startswith('._') or '__MACOSX' in info.filename:
                        continue
                    
                    ext = Path(info.filename).suffix.lower()
                    if ext not in self.SUPPORTED_TYPES:
                        continue
                    
                    try:
                        content = zf.read(info.filename)
                        doc = self.parse_file(content, info.filename)
                        results.append(doc)
                    except Exception as e:
                        errors.append({"filename": info.filename, "error": str(e)})
        
        except zipfile.BadZipFile:
            errors.append({"filename": "archive", "error": "Invalid zip file"})
        
        return results, errors
    
    def _parse_text(self, content: bytes, filename: str) -> Tuple[str, Dict]:
        """Parse plain text files"""
        if not content:
            return "", {
                'encoding': 'utf-8',
                'lines': 0,
                'characters': 0,
                'words': 0,
            }
        
        # Try different encodings
        encodings = ['utf-8', 'utf-16', 'ascii', 'iso-8859-1', 'cp1252']
        text = None
        used_encoding = 'utf-8'
        
        for encoding in encodings:
            try:
                text = content.decode(encoding)
                used_encoding = encoding
                break
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        if text is None:
            text = content.decode('utf-8', errors='ignore')
            used_encoding = 'utf-8 (fallback with ignore)'
        
        # Clean up text
        text = self._clean_text(text)
        
        metadata = {
            'encoding': used_encoding,
            'lines': text.count('\n') + 1,
            'characters': len(text),
            'words': len(text.split()),
        }
        
        return text, metadata
    
    def _parse_html(self, content: bytes, filename: str) -> Tuple[str, Dict]:
        """Parse HTML files and extract text from body tags only"""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            # Fallback to regex-based stripping
            return self._parse_html_fallback(content, filename)
        
        text, meta = self._parse_text(content, filename)
        soup = BeautifulSoup(text, 'html.parser')
        
        # Extract title from the whole document (metadata)
        title = None
        title_tag = soup.find('title')
        if title_tag:
            title = title_tag.get_text().strip()
        h1 = soup.find('h1')
        if not title and h1:
            title = h1.get_text().strip()
        
        # Extract text content from body only
        body = soup.find('body')
        if body:
            for script in body(["script", "style", "nav", "footer", "header"]):
                script.decompose()
            text = body.get_text()
        else:
            # Fallback: use whole document if no body tag
            for script in soup(["script", "style", "nav", "footer", "header"]):
                script.decompose()
            text = soup.get_text()
        
        text = self._clean_text(text)
        
        metadata = {
            **meta,
            'title': title,
            'has_tables': bool(soup.find_all('table')),
            'has_links': bool(soup.find_all('a')),
        }
        
        return text, metadata
    
    def _parse_html_fallback(self, content: bytes, filename: str) -> Tuple[str, Dict]:
        """Fallback HTML parsing without BeautifulSoup"""
        text, meta = self._parse_text(content, filename)
        
        # Extract body content only
        match = re.search(r'<body[^>]*>(.*?)</body>', text, re.DOTALL | re.IGNORECASE)
        if match:
            text = match.group(1)
        
        # Simple regex-based tag removal
        text = re.sub(r'<script[^>]*>.*?</script>', '', text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r'<[^>]+>', ' ', text)
        text = self._clean_text(text)
        
        return text, meta
    
    def _parse_pdf(self, content: bytes, filename: str) -> Tuple[str, Dict]:
        """Parse PDF files"""
        try:
            import PyPDF2
        except ImportError:
            raise ImportError("PyPDF2 not installed. Install with: pip install PyPDF2")
        
        text_parts = []
        total_pages = 0
        
        try:
            with io.BytesIO(content) as pdf_file:
                reader = PyPDF2.PdfReader(pdf_file)
                total_pages = len(reader.pages)
                
                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(f"\n--- Page {page_num + 1} ---\n")
                            text_parts.append(page_text)
                    except Exception as e:
                        text_parts.append(f"\n[Error extracting page {page_num + 1}: {e}]\n")
        
        except Exception as e:
            raise ValueError(f"Could not parse PDF: {e}")
        
        text = '\n'.join(text_parts)
        text = self._clean_text(text)
        
        metadata = {
            'pages': total_pages,
            'characters': len(text),
            'words': len(text.split()),
        }
        
        return text, metadata
    
    def _parse_docx(self, content: bytes, filename: str) -> Tuple[str, Dict]:
        """Parse DOCX files"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        
        try:
            with io.BytesIO(content) as docx_file:
                doc = Document(docx_file)
                
                text_parts = []
                
                # Extract paragraphs
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)
                
                # Extract tables
                for table_idx, table in enumerate(doc.tables):
                    text_parts.append(f"\n--- Table {table_idx + 1} ---")
                    for row in table.rows:
                        row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                        text_parts.append(row_text)
                
                text = '\n'.join(text_parts)
                text = self._clean_text(text)
                
                metadata = {
                    'paragraphs': len(doc.paragraphs),
                    'tables': len(doc.tables),
                    'characters': len(text),
                    'words': len(text.split()),
                }
                
                return text, metadata
        
        except Exception as e:
            raise ValueError(f"Could not parse DOCX: {e}")
    
    def _parse_doc(self, content: bytes, filename: str) -> Tuple[str, Dict]:
        """Parse old DOC format files"""
        # Try antiword if available
        try:
            with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
                tmp.write(content)
                tmp_path = tmp.name
            
            import subprocess
            result = subprocess.run(
                ['antiword', tmp_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            
            os.unlink(tmp_path)
            
            if result.returncode == 0:
                text = self._clean_text(result.stdout)
                return text, {'method': 'antiword', 'words': len(text.split())}
        
        except (subprocess.TimeoutExpired, FileNotFoundError):
            pass
        
        # Fallback: try to extract text directly (may contain binary garbage)
        text, meta = self._parse_text(content, filename)
        # Remove binary characters
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f]', '', text)
        
        return text, {**meta, 'method': 'fallback', 'note': 'Binary content may be present'}
    
    def _parse_xlsx(self, content: bytes, filename: str) -> Tuple[str, Dict]:
        """Parse Excel files"""
        try:
            import openpyxl
        except ImportError:
            raise ImportError("openpyxl not installed. Install with: pip install openpyxl")
        
        try:
            with io.BytesIO(content) as xlsx_file:
                wb = openpyxl.load_workbook(xlsx_file, data_only=True)
                
                text_parts = []
                total_sheets = len(wb.sheetnames)
                
                for sheet_name in wb.sheetnames:
                    sheet = wb[sheet_name]
                    text_parts.append(f"\n=== Sheet: {sheet_name} ===")
                    
                    for row in sheet.iter_rows(values_only=True):
                        # Skip empty rows
                        if any(cell is not None and str(cell).strip() for cell in row):
                            row_text = ' | '.join(
                                str(cell) if cell is not None else '' 
                                for cell in row
                            )
                            text_parts.append(row_text)
                
                text = '\n'.join(text_parts)
                text = self._clean_text(text)
                
                metadata = {
                    'sheets': total_sheets,
                    'sheet_names': wb.sheetnames,
                    'characters': len(text),
                    'rows': len(text.split('\n')),
                }
                
                return text, metadata
        
        except Exception as e:
            raise ValueError(f"Could not parse Excel file: {e}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove excessive whitespace
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Remove null bytes
        text = text.replace('\x00', '')
        
        return text.strip()
    
    def get_file_info(self, filename: str) -> Dict:
        """Get information about a file type"""
        ext = Path(filename).suffix.lower()
        return {
            'extension': ext,
            'mime_type': self.SUPPORTED_TYPES.get(ext, 'unknown'),
            'supported': ext in self.SUPPORTED_TYPES
        }


# Convenience function
def parse_uploaded_file(file_content: bytes, filename: str) -> ParsedDocument:
    """Parse a single uploaded file"""
    parser = FileParser()
    return parser.parse_file(file_content, filename)


def parse_uploaded_files(files: List[Tuple[bytes, str]]) -> Tuple[List[ParsedDocument], List[Dict]]:
    """Parse multiple uploaded files"""
    parser = FileParser()
    return parser.parse_multiple(files)


def parse_zip_archive(zip_content: bytes) -> Tuple[List[ParsedDocument], List[Dict]]:
    """Parse a zip archive"""
    parser = FileParser()
    return parser.parse_zip(zip_content)
