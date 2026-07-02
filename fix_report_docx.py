from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

src = r'C:\Users\megab\Downloads\aegis_analysis_20260701_0052_corrected.docx'
dst = r'C:\Users\megab\Downloads\aegis_analysis_20260701_0052_corrected.docx'

doc = Document(src)


def insert_after(paragraph, text, style='Normal'):
    new_p = OxmlElement('w:p')
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.text = text
    if style:
        new_para.style = style
    return new_para


# --- Section 2: align elements with s 231 burglary wording ---
doc.paragraphs[31].text = "Elements the Prosecution Must Prove:"
doc.paragraphs[32].text = "1. The defendant entered any building or ship without authority."
doc.paragraphs[33].text = "2. The defendant entered with intent to commit any imprisonable offence in that building or ship."

# --- Section 3: fix bullet characters ---
for idx in range(37, 49):
    p = doc.paragraphs[idx]
    if p.text.startswith('�'):
        p.text = '•' + p.text[1:]

# --- Section 4: Assessment of Prosecution Case ---
doc.paragraphs[50].text = (
    "1. Fingerprint and possession evidence: A fingerprint matching Harper was recovered from the inside surface of the rear service door, "
    "and Harper was arrested nearby in possession of a crowbar, lock-picking kit, an Apple MacBook Pro and a Samsung Galaxy Tablet."
)
doc.paragraphs[51].text = (
    "2. CCTV footage: The footage shows an individual entering TechHub Electronics at 23:47:11 and leaving at 23:58:11 with a noticeably fuller backpack. "
    "The Crown will rely on the timing and Harper's presence in the area as circumstantial proof of participation."
)
doc.paragraphs[52].text = "Weaknesses"
doc.paragraphs[53].text = (
    "3. CCTV identification: The footage does not clearly identify the individual as Harper. The images may be grainy, poorly lit or fail to show facial features; "
    "the clothing description is not definitive and may conflict with other witness accounts."
)
doc.paragraphs[54].text = (
    "4. Fingerprint timing and context: The fingerprint on the rear service door does not establish when Harper touched it. He may have had a lawful reason to be at the door on an earlier occasion, "
    "or the print may have been deposited after the alleged burglary. The Crown cannot prove it was left during the offending."
)
doc.paragraphs[55].text = (
    "5. Inventory and serial-number discrepancies: Rachel Thompson identified serial numbers that appear both as sold/returned stock and as stolen property. "
    "Unless the Crown can exclude the possibility that the recovered devices were legitimately sourced, possession does not prove theft."
)
doc.paragraphs[56].text = (
    "6. Witness inconsistencies: Michael Roberts' account of seeing a person near the premises at about 11:35 pm conflicts with the CCTV timings and other witness descriptions of clothing and alarm activation. "
    "These inconsistencies undermine the reliability of any identification and the Crown's narrative."
)

# --- Section 5: Evidence Analysis (replace and expand) ---
doc.paragraphs[58].text = "Fingerprint evidence"
doc.paragraphs[58].style = 'Heading 2'
doc.paragraphs[59].text = (
    "Senior Fingerprint Officer Nathan Brooks reports a fingerprint matching Harper on the inside surface of the rear service door. "
    "The defence should obtain the full comparison notes, photographs and chain-of-custody records for the lift. "
    "The critical issue is timing: a fingerprint is a static item of evidence that cannot be dated by sight. "
    "If Harper had ever attended the premises lawfully — for example as a customer, delivery person or contractor — the print could have been deposited well before 15 April 2026. "
    "The Crown must also exclude contamination during collection, packaging or transport. Without evidence that the print was left during the burglary, its probative value is limited and its prejudicial effect high."
)
p = insert_after(doc.paragraphs[59], "CCTV footage", "Heading 2")
p = insert_after(p, (
    "The CCTV footage is capable of showing that someone entered and left TechHub Electronics at the relevant times, but it does not, on its own, identify that person as Harper. "
    "The defence should obtain the original recordings, metadata, camera maintenance logs and any exported copies. "
    "Issues to explore include image resolution, lighting, frame rate, whether the timestamp is accurate, and whether the cameras cover all entry/exit points. "
    "If the footage is ambiguous, the jury should be warned that identification from it is unsafe."
), "Normal")
p = insert_after(p, "Possession of seized items", "Heading 2")
p = insert_after(p, (
    "Harper's possession of the MacBook Pro (Serial MP442781), Samsung Galaxy Tablet (Serial GT889123), crowbar and lock-picking kit is relied on as post-offence conduct. "
    "Possession of recently stolen property may support an inference of theft, but only if the Crown proves the items were stolen. "
    "Rachel Thompson's evidence that some serial numbers appear both sold and stolen directly undermines that inference. "
    "The defence should require the Crown to prove each item's provenance and to exclude innocent explanations such as purchase, gift or second-hand acquisition."
), "Normal")
p = insert_after(p, "Witness statements", "Heading 2")
p = insert_after(p, (
    "Michael Roberts says he saw a person near TechHub Electronics at about 11:35 pm. That time is materially earlier than the CCTV entry time of 23:47:11 and may relate to a different individual. "
    "The defence should test his description, lighting conditions, distance, opportunity to observe, and whether his account has changed between statements. "
    "Rachel Thompson's evidence is primarily documentary; she should be asked about the inventory system's reliability, who had access to alter records, and any communications with Mr Chen about the disputed serial numbers."
), "Normal")
p = insert_after(p, "Chain of custody and forensic gaps", "Heading 2")
insert_after(p, (
    "The integrity of the seized exhibits depends on a complete chain of custody. The defence should scrutinise Constable Patel's notebook, body-worn camera footage, exhibit packaging, storage records and any transfers to the fingerprint laboratory. "
    "The DNA sample from the damaged lock area is still pending; if it excludes Harper, that weakens the Crown case. If the chain is broken, an application to exclude evidence under the Evidence Act 2006 or NZBORA s 21 should be considered."
), "Normal")

# --- Section 6: Elements (rewrite for burglary) ---
doc.paragraphs[61].text = "1. The defendant entered any building or ship without authority."
doc.paragraphs[62].text = "Prosecution Evidence: CCTV footage shows a person entering TechHub Electronics through the rear service door at 23:47:11 on 15 April 2026; a fingerprint matching Harper was found on the inside surface of that door."
doc.paragraphs[63].text = "Defence Response/Weaknesses: The CCTV does not identify Harper. The fingerprint cannot be dated and may have been deposited lawfully on an earlier occasion. The Crown must prove Harper was the person who entered and that he did so without authority."
doc.paragraphs[64].text = "Assessment: UNCLEAR beyond reasonable doubt."
doc.paragraphs[65].text = "2. The defendant entered with intent to commit any imprisonable offence in that building or ship."
doc.paragraphs[66].text = "Prosecution Evidence: The person left 11 minutes later with a fuller backpack; Harper was arrested nearby with electronic devices and burglary tools; the rear door showed signs of forced entry."
doc.paragraphs[67].text = "Defence Response/Weaknesses: The devices in Harper's possession may not be stolen because of inventory discrepancies. The burglary tools are lawful to own and do not prove intent at the time of entry. The Crown must infer intent from circumstances, and the defence can point to innocent explanations."
doc.paragraphs[68].text = "Assessment: UNCLEAR beyond reasonable doubt."

# --- Section 7: Defence Strategies (replace and expand) ---
doc.paragraphs[70].text = "1. Challenge the admissibility of the search and seizure"
doc.paragraphs[71].text = "• Require the Crown to prove Constable Patel had reasonable suspicion under s 18(2) Search and Surveillance Act 2012 before searching Harper's bag."
doc.paragraphs[72].text = "• Obtain the officer's notebook, body-worn camera footage and any radio communications."
doc.paragraphs[73].text = "• If the search was unlawful, apply for exclusion of the seized items under NZBORA s 21 and s 30 Evidence Act 2006."
doc.paragraphs[74].text = "2. Challenge identification evidence"

p = insert_after(doc.paragraphs[74], "• Test whether CCTV quality is sufficient to identify Harper; obtain original footage, metadata and maintenance logs.", "Normal")
p = insert_after(p, "• Cross-examine Michael Roberts on lighting, distance, time and any deviations between his statements.", "Normal")
p = insert_after(p, "• Request a jury direction on the risks of identification evidence if the Crown relies on it.", "Normal")

p = insert_after(p, "3. Challenge the fingerprint evidence", "Normal")
p = insert_after(p, "• Obtain full comparison notes, photographs and chain-of-custody records from Senior Fingerprint Officer Nathan Brooks.", "Normal")
p = insert_after(p, "• Explore whether Harper had any lawful reason to touch the rear service door before the alleged offence.", "Normal")
p = insert_after(p, "• Instruct an independent fingerprint expert if the comparison or methodology is questionable.", "Normal")

p = insert_after(p, "4. Challenge proof that the seized items were stolen", "Normal")
p = insert_after(p, "• Require the Crown to prove each serial number was stolen and not sold, returned or legitimately acquired.", "Normal")
p = insert_after(p, "• Cross-examine Rachel Thompson on inventory system reliability and communications with Mr Chen.", "Normal")
p = insert_after(p, "• Obtain purchase records, receipts or supplier documentation for the recovered devices.", "Normal")

p = insert_after(p, "5. Exploit chain-of-custody and forensic gaps", "Normal")
p = insert_after(p, "• Request complete exhibit handling records, packaging details and laboratory transfer logs.", "Normal")
p = insert_after(p, "• Obtain pending DNA results from the damaged lock area; if they exclude Harper, emphasise that absence.", "Normal")
p = insert_after(p, "• Apply to exclude or discount any exhibit whose provenance cannot be proved.", "Normal")

p = insert_after(p, "6. Develop an innocent explanation for Harper's presence and possession", "Normal")
p = insert_after(p, "• Take full instructions on why Harper was in the Upper Queen Street area and whether he had a lawful explanation for the items in his bag.", "Normal")
p = insert_after(p, "• Gather corroborating witnesses, receipts or digital evidence (phone location, messages, transactions) that support that explanation.", "Normal")
insert_after(p, "• If credible, advance that explanation in evidence or through counsel's cross-examination.", "Normal")

# --- Section 8: Cross-Examination Priorities (replace and expand) ---
doc.paragraphs[76].text = "Constable Sarah Patel (seizing officer)"
doc.paragraphs[77].text = "• What was the basis for your suspicion before searching Harper's bag under s 18(2) Search and Surveillance Act 2012?"
doc.paragraphs[78].text = "• Did you make a contemporaneous record of every item seized, and does that record match the exhibit list?"
doc.paragraphs[79].text = "• What steps did you take to maintain the chain of custody for each exhibit between seizure and transfer to the fingerprint laboratory?"

p = insert_after(doc.paragraphs[79], "Why it matters: If the search was unlawful or the exhibits were mishandled, the defence can apply to exclude the items and any fingerprint evidence derived from them.", "Normal")
p = insert_after(p, "Where the defence goes next: Obtain the notebook, body-worn footage and exhibit handling records; compare them with the Crown's exhibit list.", "Normal")

p = insert_after(p, "Senior Fingerprint Officer Nathan Brooks", "Heading 2")
p = insert_after(p, "• Can you date the fingerprint or exclude the possibility that it was deposited before 15 April 2026?", "Normal")
p = insert_after(p, "• What methodology did you use to compare the lifted print with Harper's known prints, and what is the statistical basis for your conclusion?", "Normal")
p = insert_after(p, "• Were you provided with the lift, photographs and chain-of-custody documents in unbroken packaging?", "Normal")
p = insert_after(p, "Why it matters: A static fingerprint without timing or reliable methodology cannot safely prove Harper entered as a trespasser during the burglary.", "Normal")
p = insert_after(p, "Where the defence goes next: Instruct an independent expert to review the comparison notes and consider a s 25 Evidence Act challenge if the methodology is not sufficiently reliable.", "Normal")

p = insert_after(p, "Rachel Thompson (inventory/staff witness)", "Heading 2")
p = insert_after(p, "• Which specific serial numbers appear both as sold/returned stock and as stolen property, and what is the source of each entry?", "Normal")
p = insert_after(p, "• Who has access to amend inventory records, and are there audit logs showing when entries were changed?", "Normal")
p = insert_after(p, "• Did you or Mr Chen communicate about these discrepancies before or after the burglary?", "Normal")
p = insert_after(p, "Why it matters: If recovered devices were legitimately sold or returned, possession does not prove theft and the Crown's narrative collapses.", "Normal")
p = insert_after(p, "Where the defence goes next: Obtain the full inventory database, audit logs and any communications between Thompson and Chen.", "Normal")

p = insert_after(p, "Michael Roberts (eyewitness)", "Heading 2")
p = insert_after(p, "• At what exact time did you see the person, how far away were you, and what were the lighting conditions?", "Normal")
p = insert_after(p, "• Can you describe the person's clothing, height, build and any distinctive features with certainty?", "Normal")
p = insert_after(p, "• Have you given more than one statement, and has your description or timing changed?", "Normal")
p = insert_after(p, "Why it matters: Roberts' reported time of 11:35 pm is before the CCTV entry time and may describe an innocent passer-by; inconsistencies damage his reliability.", "Normal")
p = insert_after(p, "Where the defence goes next: Compare Roberts' statements with the CCTV timestamps and any other witness descriptions; prepare a timeline to show inconsistencies.", "Normal")

p = insert_after(p, "CCTV custodian/operator", "Heading 2")
p = insert_after(p, "• Are the cameras motion-activated or continuously recorded, and what is the frame rate and resolution?", "Normal")
p = insert_after(p, "• Has the recording system been maintained, and are the timestamps synchronised to an accurate clock?", "Normal")
p = insert_after(p, "• Provide the original export logs and any copies made for police.", "Normal")
p = insert_after(p, "Why it matters: Poor-quality or tampered footage cannot support a safe identification, and gaps in coverage may create reasonable doubt.", "Normal")
insert_after(p, "Where the defence goes next: Obtain original files, maintenance logs and an independent video analyst if necessary.", "Normal")

# --- Section 9: Disclosure and Forensic Gaps (add items) ---
p = insert_after(doc.paragraphs[86], "3. DNA results from the damaged lock area", "Normal")
p = insert_after(p, "Why it Matters: If the DNA excludes Harper, it weakens the Crown's claim that he forced entry; if it matches another person, that raises the possibility of an alternative offender.", "Normal")
p = insert_after(p, "Request/Application: Obtain the DNA analysis report, sample provenance and laboratory notes as soon as available.", "Normal")
p = insert_after(p, "4. CCTV original recordings and metadata", "Normal")
p = insert_after(p, "Why it Matters: Exported or compressed copies may lose detail; original files and metadata are needed to verify timestamps, integrity and authenticity.", "Normal")
insert_after(p, "Request/Application: Request the original CCTV files, system logs, camera maintenance records and a written statement from the custodian.", "Normal")

# --- Section 11: Evidentiary Issues (rewrite and expand) ---
doc.paragraphs[106].text = "1. Reliability and voluntariness of any alleged admission"
doc.paragraphs[107].text = (
    "If the Crown seeks to rely on any statement allegedly made by Harper, the defence should challenge its voluntariness and reliability. "
    "NZBORA s 24(1)(d) provides that 'everyone who is arrested or detained has the right ... to refrain from making any statement and to be informed of that right.' "
    "A breach may trigger exclusion under s 30 Evidence Act 2006."
)
doc.paragraphs[108].text = "Quotation: NZBORA s 24(1)(d) — 'Everyone who is arrested or detained has the right ... to refrain from making any statement and to be informed of that right.'"

p = insert_after(doc.paragraphs[108], "2. Admissibility of identification evidence", "Normal")
p = insert_after(p, "The CCTV footage and any eyewitness identification must be sufficiently reliable. The defence should seek a jury direction on the risks of mistaken identification and, if the evidence is weak, argue it is inadmissible or insufficient to prove identity.", "Normal")
p = insert_after(p, "3. Admissibility of fingerprint evidence", "Normal")
p = insert_after(p, "Fingerprint evidence is expert opinion evidence. The defence can challenge admissibility under s 25 Evidence Act 2006 if the methodology is not sufficiently reliable, or under s 8 if the probative value is outweighed by prejudicial effect, particularly where the print cannot be dated.", "Normal")
p = insert_after(p, "4. Admissibility of items seized from Harper", "Normal")
p = insert_after(p, "If the search was unlawful or conducted without reasonable suspicion under s 18(2) Search and Surveillance Act 2012, the seized items may be excluded under NZBORA s 21 and s 30 Evidence Act 2006.", "Normal")
p = insert_after(p, "5. Hearsay and reliability of business records", "Normal")
insert_after(p, "Inventory records are hearsay unless admitted under a hearsay exception or as business records under s 19 Evidence Act 2006. The defence should require proof of the records' accuracy, the reliability of the system, and audit trails showing who created or modified entries.", "Normal")

# --- Section 12: Conclusion (rewrite) ---
doc.paragraphs[110].text = (
    "The prosecution case against Jordan Harper is circumstantial. Its strength depends on whether the jury accepts that the fingerprint was deposited during the burglary, "
    "that the person in the CCTV footage is Harper, and that the items in his possession were stolen. Each of those propositions is contested."
)
doc.paragraphs[111].text = (
    "If the defence can show that the fingerprint could have been left innocently, that the CCTV is unclear or its timing unreliable, that the inventory records undermine proof of theft, "
    "or that the search was unlawful, the Crown will struggle to prove burglary beyond reasonable doubt. The defence's priority is to secure disclosure of chain-of-custody records, "
    "body-worn camera footage, full fingerprint comparison notes, inventory audit logs and CCTV originals. Cross-examination should focus on identification, timing and the integrity of the exhibits."
)
insert_after(doc.paragraphs[111], (
    "Conditional outcomes: If the fingerprint and CCTV evidence remain unchallenged, the Crown has a workable but not overwhelming case. "
    "If any of the key weaknesses are established, the case becomes significantly weaker and a not-guilty verdict or a successful exclusion application is realistic."
), "Normal")

doc.save(dst)
print(f"Saved fully corrected document to: {dst}")
