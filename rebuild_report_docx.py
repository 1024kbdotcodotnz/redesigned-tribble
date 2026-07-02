from docx import Document
from docx.shared import Pt

src = r'C:\Users\megab\Downloads\aegis_analysis_20260701_0052.docx'
dst = r'C:\Users\megab\Downloads\aegis_analysis_20260701_0052_corrected.docx'

doc = Document(src)

# Keep title block (paragraphs 0-6); remove everything after it.
body = doc.element.body
for p in list(doc.paragraphs)[7:]:
    body.remove(p._element)


def add(text, style='Normal'):
    p = doc.add_paragraph(text, style=style)
    return p


# Add blank line before TOC
add("")
add("TABLE OF CONTENTS", "Heading 1")
add("")
add("1. EXECUTIVE SUMMARY")
add("2. CHARGE AND LEGISLATIVE FRAMEWORK")
add("3. SUMMARY OF EVIDENCE")
add("4. ASSESSMENT OF PROSECUTION CASE")
add("5. EVIDENCE ANALYSIS")
add("6. ELEMENTS THE PROSECUTION MUST PROVE")
add("7. DEFENCE STRATEGIES AND OPTIONS")
add("8. CROSS-EXAMINATION PRIORITIES")
add("9. DISCLOSURE AND FORENSIC GAPS")
add("10. INSTRUCTIONS TO COUNSEL PRE-TRIAL")
add("11. EVIDENTIARY ISSUES TO RAISE")
add("12. CONCLUSION")
add("")

# 1. EXECUTIVE SUMMARY
add("1. EXECUTIVE SUMMARY", "Heading 1")
add(
    "On 15 April 2026, at approximately 23:47:11, a person entered TechHub Electronics through the "
    "rear service door. The same person is recorded leaving approximately 11 minutes later, at "
    "23:58:11, carrying a noticeably fuller backpack. A fingerprint matching Jordan Harper was "
    "recovered from the inside surface of the rear service door. Harper was arrested nearby on Upper "
    "Queen Street shortly after midnight (00:14 hours) and was found in possession of a MacBook Pro, "
    "a Samsung Galaxy Tablet, a crowbar and a lock-picking kit. The Crown case rests on this evidence "
    "to prove Harper's unlawful entry and his intent to commit burglary."
)
add(
    "The defence challenges the reliability and interpretation of that evidence. The CCTV footage does "
    "not clearly identify Harper; witness statements conflict on timing, clothing and location; inventory "
    "records contain serial-number discrepancies that suggest some recovered items may have been previously "
    "sold or returned; and the chain-of-custody documentation for the seized exhibits is incomplete. "
    "Rachel Thompson has identified serial numbers that appear both as sold stock and as stolen property, "
    "while Michael Roberts' account of seeing a person near the premises is inconsistent with other timings."
)
add(
    "Overall, the prosecution case is circumstantial and rests on evidence that is capable of more than "
    "one interpretation. The defence position is that Harper did not enter TechHub Electronics as a "
    "trespasser with intent to steal, that the fingerprint may have been deposited innocently at another "
    "time, that the items in his possession may not be stolen property, and that the Crown cannot exclude "
    "reasonable doubt. The immediate priorities are to obtain complete chain-of-custody records, clarify "
    "the inventory discrepancies, secure any body-worn camera footage of the search, and prepare "
    "cross-examination that exposes the identification and forensic gaps."
)

# 2. CHARGE AND LEGISLATIVE FRAMEWORK
add("2. CHARGE AND LEGISLATIVE FRAMEWORK", "Heading 1")
add("Legal Basis: Burglary under Crimes Act 1961 s 231.")
add("Relevant Legislation: Crimes Act 1961, s 231.")
add("Maximum Penalty: Not stated in the charging document.")
add("Elements the Prosecution Must Prove:")
add("1. The defendant entered any building or ship without authority.")
add("2. The defendant entered with intent to commit any imprisonable offence in that building or ship.")

# 3. SUMMARY OF EVIDENCE
add("3. SUMMARY OF EVIDENCE", "Heading 1")
add(
    "Police searched Harper near Upper Queen Street at approximately 00:14 hours on 15 April 2026, "
    "finding electronic devices and burglary tools."
)
add("Seized Evidence:", "Heading 2")
add("• Apple MacBook Pro (Serial MP442781)")
add("• Samsung Galaxy Tablet (Serial GT889123)")
add("• Crowbar")
add("• Lock Picking Kit")
add("Forensic Testing:", "Heading 2")
add("• Fingerprint evidence linking Harper to the rear service door.")
add("• DNA sample from the damaged lock area pending analysis.")
add("CCTV Footage:", "Heading 2")
add("• Shows an individual entering TechHub Electronics at 23:47:11 and leaving at 23:58:11.")
add("Witness Statements:", "Heading 2")
add("• Michael Roberts observed a person near TechHub Electronics at approximately 11:35 pm on the night of the burglary.")
add("• Rachel Thompson identified discrepancies in inventory records for some recovered property.")

# 4. ASSESSMENT OF PROSECUTION CASE
add("4. ASSESSMENT OF PROSECUTION CASE", "Heading 1")
add("Strengths", "Heading 2")
add(
    "1. Fingerprint and possession evidence: A fingerprint matching Harper was recovered from the inside "
    "surface of the rear service door, and Harper was arrested nearby in possession of a crowbar, lock-picking "
    "kit, an Apple MacBook Pro and a Samsung Galaxy Tablet."
)
add(
    "2. CCTV footage: The footage shows an individual entering TechHub Electronics at 23:47:11 and leaving at "
    "23:58:11 with a noticeably fuller backpack. The Crown will rely on the timing and Harper's presence in the "
    "area as circumstantial proof of participation."
)
add("Weaknesses", "Heading 2")
add(
    "3. CCTV identification: The footage does not clearly identify the individual as Harper. The images may be "
    "grainy, poorly lit or fail to show facial features; the clothing description is not definitive and may conflict "
    "with other witness accounts."
)
add(
    "4. Fingerprint timing and context: The fingerprint on the rear service door does not establish when Harper "
    "touched it. He may have had a lawful reason to be at the door on an earlier occasion, or the print may have been "
    "deposited after the alleged burglary. The Crown cannot prove it was left during the offending."
)
add(
    "5. Inventory and serial-number discrepancies: Rachel Thompson identified serial numbers that appear both as "
    "sold/returned stock and as stolen property. Unless the Crown can exclude the possibility that the recovered "
    "devices were legitimately sourced, possession does not prove theft."
)
add(
    "6. Witness inconsistencies: Michael Roberts' account of seeing a person near the premises at about 11:35 pm "
    "conflicts with the CCTV timings and other witness descriptions of clothing and alarm activation. These "
    "inconsistencies undermine the reliability of any identification and the Crown's narrative."
)

# 5. EVIDENCE ANALYSIS
add("5. EVIDENCE ANALYSIS", "Heading 1")
add("Fingerprint evidence", "Heading 2")
add(
    "Senior Fingerprint Officer Nathan Brooks reports a fingerprint matching Harper on the inside surface of the "
    "rear service door. The defence should obtain the full comparison notes, photographs and chain-of-custody records "
    "for the lift. The critical issue is timing: a fingerprint is a static item of evidence that cannot be dated by sight. "
    "If Harper had ever attended the premises lawfully — for example as a customer, delivery person or contractor — "
    "the print could have been deposited well before 15 April 2026. The Crown must also exclude contamination during "
    "collection, packaging or transport. Without evidence that the print was left during the burglary, its probative value "
    "is limited and its prejudicial effect high."
)
add("CCTV footage", "Heading 2")
add(
    "The CCTV footage is capable of showing that someone entered and left TechHub Electronics at the relevant times, "
    "but it does not, on its own, identify that person as Harper. The defence should obtain the original recordings, "
    "metadata, camera maintenance logs and any exported copies. Issues to explore include image resolution, lighting, "
    "frame rate, whether the timestamp is accurate, and whether the cameras cover all entry/exit points. If the footage is "
    "ambiguous, the jury should be warned that identification from it is unsafe."
)
add("Possession of seized items", "Heading 2")
add(
    "Harper's possession of the MacBook Pro (Serial MP442781), Samsung Galaxy Tablet (Serial GT889123), crowbar and "
    "lock-picking kit is relied on as post-offence conduct. Possession of recently stolen property may support an inference "
    "of theft, but only if the Crown proves the items were stolen. Rachel Thompson's evidence that some serial numbers "
    "appear both sold and stolen directly undermines that inference. The defence should require the Crown to prove each "
    "item's provenance and to exclude innocent explanations such as purchase, gift or second-hand acquisition."
)
add("Witness statements", "Heading 2")
add(
    "Michael Roberts says he saw a person near TechHub Electronics at about 11:35 pm. That time is materially earlier than "
    "the CCTV entry time of 23:47:11 and may relate to a different individual. The defence should test his description, "
    "lighting conditions, distance, opportunity to observe, and whether his account has changed between statements. "
    "Rachel Thompson's evidence is primarily documentary; she should be asked about the inventory system's reliability, "
    "who had access to alter records, and any communications with Mr Chen about the disputed serial numbers."
)
add("Chain of custody and forensic gaps", "Heading 2")
add(
    "The integrity of the seized exhibits depends on a complete chain of custody. The defence should scrutinise Constable "
    "Patel's notebook, body-worn camera footage, exhibit packaging, storage records and any transfers to the fingerprint "
    "laboratory. The DNA sample from the damaged lock area is still pending; if it excludes Harper, that weakens the Crown "
    "case. If the chain is broken, an application to exclude evidence under the Evidence Act 2006 or NZBORA s 21 should be "
    "considered."
)

# 6. ELEMENTS THE PROSECUTION MUST PROVE
add("6. ELEMENTS THE PROSECUTION MUST PROVE", "Heading 1")
add("1. The defendant entered any building or ship without authority.")
add(
    "Prosecution Evidence: CCTV footage shows a person entering TechHub Electronics through the rear service door at "
    "23:47:11 on 15 April 2026; a fingerprint matching Harper was found on the inside surface of that door."
)
add(
    "Defence Response/Weaknesses: The CCTV does not identify Harper. The fingerprint cannot be dated and may have been "
    "deposited lawfully on an earlier occasion. The Crown must prove Harper was the person who entered and that he did so "
    "without authority."
)
add("Assessment: UNCLEAR beyond reasonable doubt.")
add("2. The defendant entered with intent to commit any imprisonable offence in that building or ship.")
add(
    "Prosecution Evidence: The person left 11 minutes later with a fuller backpack; Harper was arrested nearby with "
    "electronic devices and burglary tools; the rear door showed signs of forced entry."
)
add(
    "Defence Response/Weaknesses: The devices in Harper's possession may not be stolen because of inventory discrepancies. "
    "The burglary tools are lawful to own and do not prove intent at the time of entry. The Crown must infer intent from "
    "circumstances, and the defence can point to innocent explanations."
)
add("Assessment: UNCLEAR beyond reasonable doubt.")

# 7. DEFENCE STRATEGIES AND OPTIONS
add("7. DEFENCE STRATEGIES AND OPTIONS", "Heading 1")
add("1. Challenge the admissibility of the search and seizure", "Heading 2")
add("• Require the Crown to prove Constable Patel had reasonable suspicion under s 18(2) Search and Surveillance Act 2012 before searching Harper's bag.")
add("• Obtain the officer's notebook, body-worn camera footage and any radio communications.")
add("• If the search was unlawful, apply for exclusion of the seized items under NZBORA s 21 and s 30 Evidence Act 2006.")

add("2. Challenge identification evidence", "Heading 2")
add("• Test whether CCTV quality is sufficient to identify Harper; obtain original footage, metadata and maintenance logs.")
add("• Cross-examine Michael Roberts on lighting, distance, time and any deviations between his statements.")
add("• Request a jury direction on the risks of identification evidence if the Crown relies on it.")

add("3. Challenge the fingerprint evidence", "Heading 2")
add("• Obtain full comparison notes, photographs and chain-of-custody records from Senior Fingerprint Officer Nathan Brooks.")
add("• Explore whether Harper had any lawful reason to touch the rear service door before the alleged offence.")
add("• Instruct an independent fingerprint expert if the comparison or methodology is questionable.")

add("4. Challenge proof that the seized items were stolen", "Heading 2")
add("• Require the Crown to prove each serial number was stolen and not sold, returned or legitimately acquired.")
add("• Cross-examine Rachel Thompson on inventory system reliability and communications with Mr Chen.")
add("• Obtain purchase records, receipts or supplier documentation for the recovered devices.")

add("5. Exploit chain-of-custody and forensic gaps", "Heading 2")
add("• Request complete exhibit handling records, packaging details and laboratory transfer logs.")
add("• Obtain pending DNA results from the damaged lock area; if they exclude Harper, emphasise that absence.")
add("• Apply to exclude or discount any exhibit whose provenance cannot be proved.")

add("6. Develop an innocent explanation for Harper's presence and possession", "Heading 2")
add("• Take full instructions on why Harper was in the Upper Queen Street area and whether he had a lawful explanation for the items in his bag.")
add("• Gather corroborating witnesses, receipts or digital evidence (phone location, messages, transactions) that support that explanation.")
add("• If credible, advance that explanation in evidence or through counsel's cross-examination.")

# 8. CROSS-EXAMINATION PRIORITIES
add("8. CROSS-EXAMINATION PRIORITIES", "Heading 1")
add("Constable Sarah Patel (seizing officer)", "Heading 2")
add("• What was the basis for your suspicion before searching Harper's bag under s 18(2) Search and Surveillance Act 2012?")
add("• Did you make a contemporaneous record of every item seized, and does that record match the exhibit list?")
add("• What steps did you take to maintain the chain of custody for each exhibit between seizure and transfer to the fingerprint laboratory?")
add("Why it matters: If the search was unlawful or the exhibits were mishandled, the defence can apply to exclude the items and any fingerprint evidence derived from them.")
add("Where the defence goes next: Obtain the notebook, body-worn footage and exhibit handling records; compare them with the Crown's exhibit list.")

add("Senior Fingerprint Officer Nathan Brooks", "Heading 2")
add("• Can you date the fingerprint or exclude the possibility that it was deposited before 15 April 2026?")
add("• What methodology did you use to compare the lifted print with Harper's known prints, and what is the statistical basis for your conclusion?")
add("• Were you provided with the lift, photographs and chain-of-custody documents in unbroken packaging?")
add("Why it matters: A static fingerprint without timing or reliable methodology cannot safely prove Harper entered as a trespasser during the burglary.")
add("Where the defence goes next: Instruct an independent expert to review the comparison notes and consider a s 25 Evidence Act challenge if the methodology is not sufficiently reliable.")

add("Rachel Thompson (inventory/staff witness)", "Heading 2")
add("• Which specific serial numbers appear both as sold/returned stock and as stolen property, and what is the source of each entry?")
add("• Who has access to amend inventory records, and are there audit logs showing when entries were changed?")
add("• Did you or Mr Chen communicate about these discrepancies before or after the burglary?")
add("Why it matters: If recovered devices were legitimately sold or returned, possession does not prove theft and the Crown's narrative collapses.")
add("Where the defence goes next: Obtain the full inventory database, audit logs and any communications between Thompson and Chen.")

add("Michael Roberts (eyewitness)", "Heading 2")
add("• At what exact time did you see the person, how far away were you, and what were the lighting conditions?")
add("• Can you describe the person's clothing, height, build and any distinctive features with certainty?")
add("• Have you given more than one statement, and has your description or timing changed?")
add("Why it matters: Roberts' reported time of 11:35 pm is before the CCTV entry time and may describe an innocent passer-by; inconsistencies damage his reliability.")
add("Where the defence goes next: Compare Roberts' statements with the CCTV timestamps and any other witness descriptions; prepare a timeline to show inconsistencies.")

add("CCTV custodian/operator", "Heading 2")
add("• Are the cameras motion-activated or continuously recorded, and what is the frame rate and resolution?")
add("• Has the recording system been maintained, and are the timestamps synchronised to an accurate clock?")
add("• Provide the original export logs and any copies made for police.")
add("Why it matters: Poor-quality or tampered footage cannot support a safe identification, and gaps in coverage may create reasonable doubt.")
add("Where the defence goes next: Obtain original files, maintenance logs and an independent video analyst if necessary.")

# 9. DISCLOSURE AND FORENSIC GAPS
add("9. DISCLOSURE AND FORENSIC GAPS", "Heading 1")
add("1. Chain-of-custody and handling records for seized exhibits")
add("Why it Matters: Prevents the Crown from proving the exhibit is the same item seized, creating a risk of contamination or substitution.")
add("Request/Application: Request detailed chain-of-custody documentation and any body-worn camera footage from Constable Sarah Patel's notebook.")
add("2. Body-worn camera footage of the search conducted by Constable Sarah Patel")
add("Why it Matters: Provides independent corroboration of the seizure process, ensuring proper handling and identification of seized items.")
add("Request/Application: Request body-worn camera footage and detailed notebook entries from Constable Sarah Patel.")
add("3. DNA results from the damaged lock area")
add("Why it Matters: If the DNA excludes Harper, it weakens the Crown's claim that he forced entry; if it matches another person, that raises the possibility of an alternative offender.")
add("Request/Application: Obtain the DNA analysis report, sample provenance and laboratory notes as soon as available.")
add("4. CCTV original recordings and metadata")
add("Why it Matters: Exported or compressed copies may lose detail; original files and metadata are needed to verify timestamps, integrity and authenticity.")
add("Request/Application: Request the original CCTV files, system logs, camera maintenance records and a written statement from the custodian.")

# 10. INSTRUCTIONS TO COUNSEL PRE-TRIAL
add("10. INSTRUCTIONS TO COUNSEL PRE-TRIAL", "Heading 1")
add("1. Obtain detailed chain-of-custody documentation for all seized exhibits.")
add("• a) Review the officer's notebook and any body-worn camera footage.")
add("• b) Identify gaps or inconsistencies in handling and storage procedures.")
add("• c) Prepare to argue exclusion under NZBORA s 21 if necessary.")
add("2. Request clarification on inventory discrepancies from Rachel Thompson.")
add("• a) Obtain detailed explanations for serial numbers appearing both sold and stolen.")
add("• b) Review any communications between Rachel Thompson and Mr Chen regarding these discrepancies.")
add("• c) Identify potential issues with the accuracy of stock records provided to police.")
add("3. Prepare cross-examination questions for key witnesses.")
add("• a) Focus on inconsistencies in witness statements, identification risks, and chain-of-custody gaps.")
add("• b) Highlight any contradictions or uncertainties that undermine the prosecution's case.")
add("4. Review CCTV footage and identify potential issues with identification.")
add("• a) Analyze lighting conditions, viewing angles, and clarity of images.")
add("• b) Identify any technical limitations or errors in the footage.")
add("5. Obtain expert analysis on fingerprint evidence reliability.")
add("• a) Request an independent review of fingerprint comparison methods used by Senior Fingerprint Officer Nathan Brooks.")
add("• b) Highlight potential contamination risks or alternative explanations for the recovered print.")

# 11. EVIDENTIARY ISSUES TO RAISE
add("11. EVIDENTIARY ISSUES TO RAISE", "Heading 1")
add("1. Reliability and voluntariness of any alleged admission")
add(
    "If the Crown seeks to rely on any statement allegedly made by Harper, the defence should challenge its voluntariness "
    "and reliability. NZBORA s 24(1)(d) provides that 'everyone who is arrested or detained has the right ... to refrain from "
    "making any statement and to be informed of that right.' A breach may trigger exclusion under s 30 Evidence Act 2006."
)
add("Quotation: NZBORA s 24(1)(d) — 'Everyone who is arrested or detained has the right ... to refrain from making any statement and to be informed of that right.'")
add("2. Admissibility of identification evidence")
add("The CCTV footage and any eyewitness identification must be sufficiently reliable. The defence should seek a jury direction on the risks of mistaken identification and, if the evidence is weak, argue it is inadmissible or insufficient to prove identity.")
add("3. Admissibility of fingerprint evidence")
add("Fingerprint evidence is expert opinion evidence. The defence can challenge admissibility under s 25 Evidence Act 2006 if the methodology is not sufficiently reliable, or under s 8 if the probative value is outweighed by prejudicial effect, particularly where the print cannot be dated.")
add("4. Admissibility of items seized from Harper")
add("If the search was unlawful or conducted without reasonable suspicion under s 18(2) Search and Surveillance Act 2012, the seized items may be excluded under NZBORA s 21 and s 30 Evidence Act 2006.")
add("5. Hearsay and reliability of business records")
add("Inventory records are hearsay unless admitted under a hearsay exception or as business records under s 19 Evidence Act 2006. The defence should require proof of the records' accuracy, the reliability of the system, and audit trails showing who created or modified entries.")

# 12. CONCLUSION
add("12. CONCLUSION", "Heading 1")
add(
    "The prosecution case against Jordan Harper is circumstantial. Its strength depends on whether the jury accepts that the "
    "fingerprint was deposited during the burglary, that the person in the CCTV footage is Harper, and that the items in his "
    "possession were stolen. Each of those propositions is contested."
)
add(
    "If the defence can show that the fingerprint could have been left innocently, that the CCTV is unclear or its timing unreliable, "
    "that the inventory records undermine proof of theft, or that the search was unlawful, the Crown will struggle to prove burglary "
    "beyond reasonable doubt. The defence's priority is to secure disclosure of chain-of-custody records, body-worn camera footage, "
    "full fingerprint comparison notes, inventory audit logs and CCTV originals. Cross-examination should focus on identification, "
    "timing and the integrity of the exhibits."
)
add(
    "Conditional outcomes: If the fingerprint and CCTV evidence remain unchallenged, the Crown has a workable but not overwhelming "
    "case. If any of the key weaknesses are established, the case becomes significantly weaker and a not-guilty verdict or a "
    "successful exclusion application is realistic."
)

add("")
add(
    "This analysis is generated by an AI model based on the provided disclosure material and legal sources. It does not constitute "
    "legal advice or representation. Always consult with a qualified lawyer for specific guidance in your jurisdiction."
)

doc.save(dst)
print(f"Rebuilt corrected document saved to: {dst}")
